import os
import warnings
import time
import shutil
import json
import re
from typing import List, Dict, Optional
from urllib.parse import unquote
from flask import Flask, render_template, request, jsonify, Response, stream_with_context, redirect, url_for

# 设置环境变量
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
os.environ['HF_HUB_CACHE'] = 'D:/api/models'  

# 忽略警告
warnings.filterwarnings("ignore")

# ===================== 导入必要的库 =====================
try:
    from langchain_community.document_loaders import TextLoader, PyMuPDFLoader
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    from langchain_core.documents import Document
    LANGCHAIN_AVAILABLE = True
except ImportError as e:
    LANGCHAIN_AVAILABLE = False
    LANGCHAIN_ERROR = str(e)
    print(f"[ERROR] 缺少必要的依赖: {e}")
    print("[INFO] 请运行以下命令安装依赖:")
    print("  pip install langchain-community langchain-huggingface langchain-text-splitters")

from openai import OpenAI

# ===================== 多知识库核心配置 =====================
# 根目录
ROOT_KNOWLEDGE_DIR = "D:/api/multi_knowledge_bases"
ROOT_VECTOR_DIR = "D:/api/multi_vector_bases"
# 允许上传的文件类型
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc', 'md', 'json'}
# 文本分割配置
CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

# 大模型配置
API_KEY = os.environ.get("GLM_API_KEY")
MODEL = "glm-4.7-flash"
BASE_URL = "https://open.bigmodel.cn/api/paas/v4/"

# 创建根目录
os.makedirs(ROOT_KNOWLEDGE_DIR, exist_ok=True)
os.makedirs(ROOT_VECTOR_DIR, exist_ok=True)

# ===================== Flask 应用初始化 =====================
app = Flask(__name__)
# 文件上传大小限制（100MB）
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

# ===================== 工具函数 =====================
def allowed_file(filename: str) -> bool:
    """检查文件类型是否允许"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_all_knowledge_bases() -> List[Dict]:
    """获取所有知识库列表"""
    kb_list = []
    if os.path.exists(ROOT_KNOWLEDGE_DIR):
        for kb_name in os.listdir(ROOT_KNOWLEDGE_DIR):
            kb_path = os.path.join(ROOT_KNOWLEDGE_DIR, kb_name)
            if os.path.isdir(kb_path):
                # 统计文件数量
                file_count = sum([len(files) for _, _, files in os.walk(kb_path)])
                # 检查是否已构建向量库
                vector_path = os.path.join(ROOT_VECTOR_DIR, kb_name)
                is_built = os.path.exists(vector_path)
                
                kb_list.append({
                    "name": kb_name,
                    "file_count": file_count,
                    "is_built": is_built,
                    "knowledge_path": kb_path,
                    "vector_path": vector_path
                })
    return kb_list

def safe_filename(name: str) -> str:
    """安全的文件名处理，保留中文字符"""
    import re
    # 移除或替换危险字符，但保留中文
    name = name.strip()
    # 替换路径分隔符和其他危险字符
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    # 限制长度
    if len(name) > 100:
        name = name[:100]
    return name


def get_kb_path(kb_name: str, type: str = "knowledge") -> str:
    """获取知识库/向量库路径"""
    safe_name = safe_filename(kb_name)
    if type == "knowledge":
        return os.path.join(ROOT_KNOWLEDGE_DIR, safe_name)
    elif type == "vector":
        return os.path.join(ROOT_VECTOR_DIR, safe_name)
    raise ValueError("type must be 'knowledge' or 'vector'")

# ===================== 核心功能函数 =====================
def load_all_files(folder_path: str) -> List[Document]:
    """加载指定文件夹下的所有支持的文件"""
    documents = []
    if not os.path.exists(folder_path):
        print(f"[DEBUG] 文件夹不存在: {folder_path}")
        return documents

    if not LANGCHAIN_AVAILABLE:
        print(f"[ERROR] LangChain依赖未安装，无法加载文件")
        return documents

    # 支持的文件扩展名
    SUPPORTED_EXTS = {'.txt', '.pdf', '.md', '.json', '.docx', '.doc'}

    # 遍历文件夹
    for root, _, files in os.walk(folder_path):
        print(f"[DEBUG] 扫描文件夹: {root}, 文件数: {len(files)}")
        for filename in files:
            file_path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()

            if ext not in SUPPORTED_EXTS:
                continue

            print(f"[DEBUG] 处理文件: {filename}, 扩展名: {ext}")

            try:
                if ext == '.pdf':
                    # PDF使用PyMuPDFLoader
                    print(f"[DEBUG] 尝试加载PDF: {filename}")
                    loader = PyMuPDFLoader(file_path)
                    docs = loader.load()
                    documents.extend(docs)
                    print(f"[DEBUG] 成功加载 PDF: {filename}, 页数: {len(docs)}")

                elif ext in ['.txt', '.md', '.json']:
                    # 纯文本类文件使用TextLoader
                    print(f"[DEBUG] 尝试加载文本: {filename}")
                    try:
                        loader = TextLoader(file_path, encoding='utf-8')
                        docs = loader.load()
                        documents.extend(docs)
                        print(f"[DEBUG] 成功加载 {ext}: {filename}")
                    except UnicodeDecodeError:
                        # UTF-8失败尝试GBK
                        loader = TextLoader(file_path, encoding='gbk')
                        docs = loader.load()
                        documents.extend(docs)
                        print(f"[DEBUG] 成功加载 {ext}(GBK): {filename}")

                elif ext in ['.docx', '.doc']:
                    # Word文档尝试使用python-docx解析
                    print(f"[DEBUG] 尝试加载Word: {filename}")
                    try:
                        import docx
                        docx_doc = docx.Document(file_path)
                        text = '\n'.join([para.text for para in docx_doc.paragraphs if para.text.strip()])
                        if text.strip():
                            documents.append(Document(
                                page_content=text,
                                metadata={"source": filename}
                            ))
                            print(f"[DEBUG] 成功加载 Word: {filename}")
                        else:
                            print(f"[DEBUG] Word文档为空: {filename}")
                    except ImportError:
                        print(f"[WARNING] 未安装python-docx，跳过Word文档: {filename}")
                    except Exception as e:
                        print(f"[ERROR] 解析Word文档 {filename} 失败: {e}")

            except Exception as e:
                print(f"[ERROR] 加载文件 {filename} 出错: {e}")
                continue

    print(f"[DEBUG] 总共加载 {len(documents)} 个文档")
    return documents


def get_kb_files(kb_name: str) -> List[Dict]:
    """获取指定知识库中的所有文件列表"""
    files = []
    kb_path = get_kb_path(kb_name, "knowledge")

    if os.path.exists(kb_path):
        for root, _, filenames in os.walk(kb_path):
            for filename in filenames:
                if '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS:
                    file_path = os.path.join(root, filename)
                    rel_path = os.path.relpath(file_path, kb_path)
                    files.append({
                        "name": filename,
                        "path": rel_path,
                        "size": os.path.getsize(file_path),
                        "modified": os.path.getmtime(file_path)
                    })

    return sorted(files, key=lambda x: x["modified"], reverse=True)

def build_knowledge_base_generator(kb_name: str):
    """生成器函数：为指定知识库构建向量库，返回进度（SSE）"""
    try:
        knowledge_path = get_kb_path(kb_name, "knowledge")
        vector_path = get_kb_path(kb_name, "vector")

        print(f"[DEBUG] 构建知识库: {kb_name}")
        print(f"[DEBUG] 知识路径: {knowledge_path}")
        print(f"[DEBUG] 向量路径: {vector_path}")

        # 步骤1：检查知识库文件夹
        yield {"progress": 5, "message": f"检查[{kb_name}]知识库文件夹..."}
        time.sleep(0.2)

        if not os.path.exists(knowledge_path):
            print(f"[ERROR] 知识库路径不存在: {knowledge_path}")
            yield {"progress": 100, "message": f"⚠️ 知识库[{kb_name}]不存在，请先创建并上传文件", "success": False}
            return

        # 步骤2：扫描并加载文件
        yield {"progress": 15, "message": f"🔍 正在扫描[{kb_name}]知识库文件"}
        time.sleep(0.2)
        
        documents = load_all_files(knowledge_path)
        yield {"progress": 30, "message": f"📂 扫描完成，找到 {len(documents)} 个文件"}
        time.sleep(0.2)

        if len(documents) == 0:
            yield {"progress": 100, "message": f"⚠️ [{kb_name}]知识库中未找到任何 TXT/PDF 文件", "success": False}
            return

        # 步骤3：分割文本
        yield {"progress": 40, "message": f"✂️ 正在分割[{kb_name}]文本片段..."}
        time.sleep(0.2)
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        splits = text_splitter.split_documents(documents)
        yield {"progress": 50, "message": f"📝 文本分割完成，共生成 {len(splits)} 个片段"}
        time.sleep(0.2)

        # 步骤4：生成嵌入向量
        yield {"progress": 60, "message": f"🔮 正在生成[{kb_name}]文本嵌入向量..."}
        time.sleep(0.2)
        
        # 使用GPU运行Embedding模型（如GPU不可用会自动使用CPU）
        embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-small-zh-v1.5",
            model_kwargs={'device': 'cuda'},
            encode_kwargs={'normalize_embeddings': True}
        )
        db = FAISS.from_documents(splits, embeddings)
        yield {"progress": 80, "message": f"✅ [{kb_name}]嵌入向量生成完成"}
        time.sleep(0.2)

        # 步骤5：保存向量库
        yield {"progress": 85, "message": f"💾 正在保存[{kb_name}]向量数据库..."}
        time.sleep(0.2)
        
        # 先删除旧向量库
        if os.path.exists(vector_path):
            shutil.rmtree(vector_path)
        db.save_local(vector_path)
        yield {"progress": 100, "message": f"🎉 [{kb_name}]知识库构建完成！", "success": True}
        
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[ERROR] 构建知识库[{kb_name}]失败:")
        print(error_detail)
        yield {"progress": 100, "message": f"❌ [{kb_name}]构建失败：{str(e)}", "success": False}

def retrieve_multi_knowledge(question: str, selected_kbs: List[str], top_k=3) -> List[Dict | str]:
    """从多个选中的知识库中检索内容"""
    all_results = []
    
    if not selected_kbs or len(selected_kbs) == 0:
        return ["⚠️ 请先选择要检索的知识库！"]
    
    embeddings = HuggingFaceEmbeddings(
        model_name="BAAI/bge-small-zh-v1.5",
        model_kwargs={'device': 'cuda'},
        encode_kwargs={'normalize_embeddings': True}
    )

    for kb_name in selected_kbs:
        vector_path = get_kb_path(kb_name, "vector")
        if not os.path.exists(vector_path):
            all_results.append(f"⚠️ [{kb_name}]知识库未构建向量库，请先构建")
            continue
        
        try:
            db = FAISS.load_local(vector_path, embeddings, allow_dangerous_deserialization=True)
            docs = db.similarity_search(question, k=top_k)
            
            for d in docs:
                all_results.append({
                    "content": d.page_content,
                    "source": f"[{kb_name}] {os.path.basename(d.metadata.get('source', '未知文件'))}"
                })
        except Exception as e:
            all_results.append(f"❌ [{kb_name}]检索失败：{str(e)}")
    
    if not all_results:
        return ["📖 选中的知识库中无相关内容"]
    return all_results

def generate_answer(question: str, selected_kbs: List[str], selected_knowledge: List[Dict] = None):
    """基于选中的知识库和知识块生成回答"""
    # 1. 从选中的知识库中检索所有内容
    all_local_knowledge = retrieve_multi_knowledge(question, selected_kbs)
    
    # 2. 确定用于推理的知识块
    if selected_knowledge and len(selected_knowledge) > 0:
        used_knowledge = [item for item in selected_knowledge if isinstance(item, dict)]
        knowledge_type = "选中的"
    else:
        used_knowledge = [item for item in all_local_knowledge if isinstance(item, dict)]
        knowledge_type = "全部"

    online_knowledge = f"✅ 已禁用联网搜索功能，使用{knowledge_type}[{','.join(selected_kbs)}]知识库内容回答"  
    
    try:
        if not API_KEY:
            final_answer = "❌ 未配置API_KEY，无法生成回答"
        else:
            # 拼接选中的知识块
            local_knowledge_str = "\n---\n".join([
                item["content"] if isinstance(item, dict) else item 
                for item in used_knowledge
            ]) if used_knowledge else "📖 选中的知识库中无相关内容"
            
            client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
            resp = client.chat.completions.create(
                model=MODEL, 
                messages=[
                    {"role":"system","content": """
你是一个专业的多知识库问答助手，回答需满足以下要求：
1. 仅基于提供的知识库内容回答，不编造信息；
2. 明确标注信息来源（知识库名称）；
3. 格式美观：使用分点（•）、编号（1.2.3.）、加粗（**内容**）突出重点；
4. 语言简洁易懂，逻辑清晰，段落分明；
5. 如果没有相关内容，友好提示"未查询到相关内容"；
6. 关键信息加粗强调，重要内容分行展示。
                    """},
                    {"role":"user","content":f"问题：{question}\n知识库内容：{local_knowledge_str}"}
                ]
            )
            final_answer = resp.choices[0].message.content
    except Exception as e:
        final_answer = f"❌ 生成回答失败：{str(e)}"
    
    return {
        "question": question,
        "local_knowledge": all_local_knowledge,
        "used_knowledge": used_knowledge,
        "online_knowledge": online_knowledge,
        "final_answer": final_answer
    }

# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """主页"""
    return render_template('index.html')  

@app.route('/api/kb/list')
def api_kb_list():
    """获取所有知识库列表"""
    try:
        kb_list = get_all_knowledge_bases()
        return jsonify({"success": True, "data": kb_list})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@app.route('/api/kb/create', methods=['POST'])
def api_kb_create():
    """创建新知识库"""
    try:
        data = request.get_json()
        kb_name = data.get('name', '').strip()
        
        if not kb_name:
            return jsonify({"success": False, "message": "知识库名称不能为空"})
        
        kb_path = get_kb_path(kb_name, "knowledge")
        if os.path.exists(kb_path):
            return jsonify({"success": False, "message": f"知识库[{kb_name}]已存在"})
        
        os.makedirs(kb_path)
        return jsonify({"success": True, "message": f"知识库[{kb_name}]创建成功"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@app.route('/api/kb/upload', methods=['POST'])
def api_kb_upload():
    """上传文件到指定知识库"""
    try:
        kb_name = request.form.get('kb_name', '').strip()
        if not kb_name:
            return jsonify({"success": False, "message": "请选择要上传的知识库"})
        
        if 'file' not in request.files:
            return jsonify({"success": False, "message": "请选择要上传的文件"})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "message": "文件名称不能为空"})
        
        if file and allowed_file(file.filename):
            kb_path = get_kb_path(kb_name, "knowledge")
            filename = safe_filename(file.filename)
            file_path = os.path.join(kb_path, filename)
            file.save(file_path)

            # 检查是否是第一个文件，如果是则自动触发构建
            files_in_kb = get_kb_files(kb_name)
            should_auto_build = len(files_in_kb) <= 1  # 只有刚上传的这个文件

            return jsonify({
                "success": True,
                "message": f"文件[{filename}]上传成功",
                "filename": filename,
                "auto_build": should_auto_build,
                "kb_name": kb_name
            })
        else:
            return jsonify({"success": False, "message": f"仅支持上传{', '.join(ALLOWED_EXTENSIONS)}文件"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@app.route('/api/kb/files', methods=['GET'])
def api_kb_files():
    """获取指定知识库的文件列表"""
    try:
        kb_name = request.args.get('name', '').strip()
        if not kb_name:
            return jsonify({"success": False, "message": "知识库名称不能为空"})

        files = get_kb_files(kb_name)
        return jsonify({"success": True, "data": files})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})


@app.route('/api/kb/delete', methods=['POST'])
def api_kb_delete():
    """删除指定知识库（含文件和向量库）"""
    try:
        data = request.get_json()
        kb_name = data.get('name', '').strip()
        
        if not kb_name:
            return jsonify({"success": False, "message": "知识库名称不能为空"})
        
        # 删除知识库文件目录
        kb_path = get_kb_path(kb_name, "knowledge")
        if os.path.exists(kb_path):
            shutil.rmtree(kb_path)
        
        # 删除向量库目录
        vector_path = get_kb_path(kb_name, "vector")
        if os.path.exists(vector_path):
            shutil.rmtree(vector_path)
        
        return jsonify({"success": True, "message": f"知识库[{kb_name}]删除成功"})
    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@app.route('/build_progress/<path:kb_name>')
def build_progress(kb_name):
    """SSE路由：返回指定知识库的构建进度"""
    kb_name = unquote(kb_name)  # URL解码中文

    @stream_with_context
    def generate():
        for item in build_knowledge_base_generator(kb_name):
            yield f"data: {json.dumps(item)}\n\n"
            if item["progress"] == 100:
                break
    return Response(generate(), mimetype='text/event-stream')

@app.route('/api/retrieve', methods=['POST'])
def api_retrieve():
    """仅检索知识块，不生成回答"""
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        selected_kbs = data.get('selected_kbs', [])

        if not question:
            return jsonify({"error": "请输入问题"}), 400

        if not selected_kbs:
            return jsonify({"error": "请选择要检索的知识库"}), 400

        # 只检索知识，不生成回答
        local_knowledge = retrieve_multi_knowledge(question, selected_kbs)

        return jsonify({
            "question": question,
            "local_knowledge": local_knowledge,
            "final_answer": None  # 不生成回答
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/chat', methods=['POST'])
def api_chat():
    """问答API：支持选择多个知识库和知识块"""
    try:
        data = request.get_json()
        question = data.get('question', '').strip()
        selected_kbs = data.get('selected_kbs', [])  # 选中的知识库列表
        selected_knowledge = data.get('selected_knowledge', [])  # 选中的知识块

        if not question:
            return jsonify({"error": "请输入问题"}), 400

        if not selected_kbs:
            return jsonify({"error": "请选择要检索的知识库"}), 400

        result = generate_answer(question, selected_kbs, selected_knowledge)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ===================== 主程序 =====================
if __name__ == "__main__":
    print("===== 多知识库问答系统 - Web版 =====")
    print("🌐 网页访问地址：http://127.0.0.1:5000")
    print("🔧 按 Ctrl+C 停止服务")
    print("======================================\n")
    
    app.run(
        host='0.0.0.0',
        port=5000,
        debug=False,
        threaded=True,
        processes=4
    )