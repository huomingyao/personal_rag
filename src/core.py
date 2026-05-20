"""核心功能模块 - RAG和问答核心逻辑"""

import os
import re
import shutil
import time
from typing import List, Dict, Optional, Generator
from urllib.parse import unquote

from langchain_community.document_loaders import TextLoader, PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document
from openai import OpenAI

from .config import (
    DEVICE, MODEL, BASE_URL, API_KEY, CHUNK_SIZE, CHUNK_OVERLAP,
    EMBEDDING_MODEL
)
from .utils import get_kb_path


# ===================== 初始化 =====================
def init_embeddings():
    """初始化Embeddings模型"""
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={'device': DEVICE},
        encode_kwargs={'normalize_embeddings': True}
    )


# 全局Embeddings实例
embeddings = init_embeddings()

# 向量库缓存
_vector_db_cache = {}


# ===================== 文件加载 =====================
def load_all_files(folder_path: str) -> List[Document]:
    """加载指定文件夹下的所有支持的文件"""
    documents = []
    if not os.path.exists(folder_path):
        return documents

    SUPPORTED_EXTS = {'.txt', '.pdf', '.md', '.json', '.docx'}

    for root, _, files in os.walk(folder_path):
        for filename in files:
            file_path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()
            if ext not in SUPPORTED_EXTS:
                continue

            try:
                if ext == '.pdf':
                    loader = PyMuPDFLoader(file_path)
                    docs = loader.load()
                    documents.extend(docs)
                elif ext in ['.txt', '.md', '.json']:
                    try:
                        loader = TextLoader(file_path, encoding='utf-8')
                        docs = loader.load()
                        documents.extend(docs)
                    except UnicodeDecodeError:
                        loader = TextLoader(file_path, encoding='gbk')
                        docs = loader.load()
                        documents.extend(docs)
                elif ext == '.docx':
                    try:
                        import docx
                        docx_doc = docx.Document(file_path)
                        text = '\n'.join([p.text for p in docx_doc.paragraphs if p.text.strip()])
                        if text.strip():
                            documents.append(Document(page_content=text, metadata={"source": filename}))
                    except ImportError:
                        pass
            except Exception as e:
                print(f"[ERROR] 加载文件 {filename} 出错: {e}")
    return documents


# ===================== Query改写 =====================
def rewrite_query(question: str) -> str:
    """将用户问题改写成更适合检索的形式"""
    if not question or not question.strip() or not API_KEY:
        return question

    try:
        client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
        resp = client.chat.completions.create(
            model=MODEL,
            messages=[
                {"role": "system", "content": """你是一个Query改写专家。
改写要求：
1. 提取关键实体（人名、机构名、术语等）
2. 添加领域同义词、近义词扩展
3. 将口语化表达改写为正式的技术描述
4. 问题过长时拆分为多个子问题（用|分隔）
5. 保持原问题��核心意图不变
输出格式：直接输出改写后的检索词，不要添加任何解释或前缀。"""},
                {"role": "user", "content": f"请改写这个问题：{question}"}
            ],
            max_tokens=500,
            temperature=0.3
        )
        rewritten = resp.choices[0].message.content
        if rewritten:
            return rewritten.strip()
    except Exception as e:
        print(f"[DEBUG] Query改写失败: {e}")
    return question


# ===================== 联网搜索 =====================
def web_search(query: str, num_results: int = 5) -> str:
    """联网搜索"""
    try:
        from duckduckgo_search import DuckDuckGoResults
        ddgs = DuckDuckGoResults()
        results = []
        for r in ddgs.text(query, max_results=num_results):
            results.append({
                "title": r.get("title", ""),
                "href": r.get("href", ""),
                "body": r.get("body", "")
            })
        if not results:
            return "未找到相关结果"

        formatted = []
        for i, r in enumerate(results, 1):
            formatted.append(f"{i}. **{r.get('title', '无标题')}**\n   {r.get('body', '无描述')}\n   来源: {r.get('href', '')}")
        return "\n\n".join(formatted)
    except ImportError:
        return "未安装duckduckgo-search库"
    except Exception as e:
        return f"搜索失败: {str(e)}"


def web_search_sync(question: str, enable_web: bool = False, num_results: int = 5) -> str:
    """联网搜索主函数"""
    if not enable_web:
        return ""
    return web_search(question, num_results)


# ===================== 知识库构建 =====================
def build_knowledge_base_generator(kb_name: str) -> Generator[Dict, None, None]:
    """生成器函数：为指定知识库构建向量库，返回进度（SSE）"""
    knowledge_path = get_kb_path(kb_name, "knowledge")
    vector_path = get_kb_path(kb_name, "vector")

    try:
        yield {"progress": 5, "message": f"检查[{kb_name}]知识库文件夹..."}
        time.sleep(0.2)

        if not os.path.exists(knowledge_path):
            yield {"progress": 100, "message": f"⚠️ 知识库[{kb_name}]不存在", "success": False}
            return

        yield {"progress": 15, "message": f"🔍 正在扫描[{kb_name}]知识库文件"}
        time.sleep(0.2)

        documents = load_all_files(knowledge_path)
        yield {"progress": 30, "message": f"📂 扫描完成，找到 {len(documents)} 个文件"}
        time.sleep(0.2)

        if len(documents) == 0:
            yield {"progress": 100, "message": f"⚠️ [{kb_name}]知识库中未找到任何文件", "success": False}
            return

        yield {"progress": 40, "message": f"✂️ 正在分割[{kb_name}]文本片段..."}
        time.sleep(0.2)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
        splits = text_splitter.split_documents(documents)

        yield {"progress": 50, "message": f"📝 文本分割完成，共生成 {len(splits)} 个片段"}
        time.sleep(0.2)

        if len(splits) == 0:
            yield {"progress": 100, "message": f"⚠️ [{kb_name}]文本分割后无有效内容", "success": False}
            return

        valid_splits = [s for s in splits if s.page_content and s.page_content.strip()]
        if len(valid_splits) < len(splits):
            print(f"[DEBUG] 过滤了 {len(splits) - len(valid_splits)} 个空文档")
            splits = valid_splits

        yield {"progress": 60, "message": f"🔮 正在生成[{kb_name}]文本嵌入向量... (设备: {DEVICE})"}
        time.sleep(0.2)

        db = FAISS.from_documents(splits, embeddings)
        yield {"progress": 80, "message": f"✅ [{kb_name}]嵌入向量生成完成"}
        time.sleep(0.2)

        yield {"progress": 85, "message": f"💾 正在保存[{kb_name}]向量数据库..."}
        time.sleep(0.2)

        if os.path.exists(vector_path):
            shutil.rmtree(vector_path)
        db.save_local(vector_path)
        yield {"progress": 100, "message": f"🎉 [{kb_name}]知识库构建完成！", "success": True}

    except Exception as e:
        import traceback
        print(f"[ERROR] 构建知识库[{kb_name}]失败:\n{traceback.format_exc()}")
        yield {"progress": 100, "message": f"❌ [{kb_name}]构建失败：{str(e)}", "success": False}


# ===================== 知识检索 =====================
def retrieve_multi_knowledge(question: str, selected_kbs: List[str], top_k: int = 3) -> tuple:
    """从多个选中的知识库中检索内容"""
    all_results = []

    if not selected_kbs:
        return ["⚠️ 请先选择要检索的知识库！"], question, question

    original_question = question
    rewritten_query = rewrite_query(question)
    if rewritten_query != question:
        print(f"[DEBUG] Query改写: '{question}' -> '{rewritten_query}'")
        question = rewritten_query

    for kb_name in selected_kbs:
        vector_path = get_kb_path(kb_name, "vector")
        if not os.path.exists(vector_path):
            all_results.append(f"⚠️ [{kb_name}]知识库未构建向量库")
            continue

        try:
            if kb_name not in _vector_db_cache:
                _vector_db_cache[kb_name] = FAISS.load_local(vector_path, embeddings, allow_dangerous_deserialization=True)
            db = _vector_db_cache[kb_name]
            docs = db.similarity_search(question, k=top_k)

            for d in docs:
                all_results.append({
                    "content": d.page_content,
                    "source": f"[{kb_name}] {os.path.basename(d.metadata.get('source', '未知文件'))}"
                })
        except Exception as e:
            all_results.append(f"❌ [{kb_name}]检索失败：{str(e)}")

    if not all_results:
        return ["📖 选中的知识库中无相关内容"], original_question, rewritten_query
    return all_results, original_question, rewritten_query


# ===================== 回答生成 =====================
def generate_answer(
    question: str,
    selected_kbs: List[str],
    selected_knowledge: List[Dict] = None,
    enable_web: bool = False
) -> Dict:
    """基于选中的知识库和知识块生成回答"""
    all_local_knowledge, original_question, rewritten_query = retrieve_multi_knowledge(question, selected_kbs)

    if selected_knowledge and len(selected_knowledge) > 0:
        used_knowledge = [item for item in selected_knowledge if isinstance(item, dict)]
        knowledge_type = "选中的"
    else:
        used_knowledge = [item for item in all_local_knowledge if isinstance(item, dict)]
        knowledge_type = "全部"

    web_knowledge = ""
    if enable_web:
        web_knowledge = web_search_sync(question, enable_web)
        if web_knowledge:
            online_knowledge = f"🌐 已启用联网搜索 + {knowledge_type}[{','.join(selected_kbs)}]知识库"
        else:
            online_knowledge = f"🌐 已启用联网搜索（无结果），使用{knowledge_type}[{','.join(selected_kbs)}]知识库"
    else:
        online_knowledge = f"✅ 已禁用联网搜索功能，使用{knowledge_type}[{','.join(selected_kbs)}]知识库内容回答"

    try:
        if not API_KEY:
            final_answer = "❌ 未配置MINIMAX_API，无法生成回答"
        else:
            local_knowledge_str = "\n---\n".join([
                item["content"] if isinstance(item, dict) else item
                for item in used_knowledge
            ]) if used_knowledge else "📖 选中的知识库中无相关内容"

            context_parts = [f"知识库内容：{local_knowledge_str}"]
            if web_knowledge:
                context_parts.append(f"联网搜索结果：{web_knowledge}")
            full_context = "\n\n".join(context_parts)

            if enable_web and web_knowledge:
                system_prompt = """你是一个专业的问答助手。回答需满足以下要求：
1. 优先使用联网搜索结果回答，同时可参考知识库内容；
2. 明确标注信息来源；
3. 格式美观：使用分点、加粗突出重点；
4. 语言简洁易懂，逻辑清晰；
5. 如果没有相关内容，友好提示"未查询到相关内容"。"""
            else:
                system_prompt = """你是一个专业的多知识库问答助手，回答需满足以下要求：
1. 仅基于提供的知识库内容回答，不编造信息；
2. 明确标注信息来源；
3. 格式美观：使用分点、加粗突出重点；
4. 语言简洁易懂，逻辑清晰；
5. 如果没有相关内容，友好提示"未查询到相关内容"。"""

            client = OpenAI(api_key=API_KEY, base_url=BASE_URL)
            resp = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"问题：{question}\n{full_context}"}
                ]
            )
            content = resp.choices[0].message.content or ""

            # 提取thinking内容
            thinking_content = ""
            if "<think>" in content and "</think>" in content:
                think_matches = re.findall(r'<think>(.*?)</think>', content, re.DOTALL)
                if think_matches:
                    thinking_content = '\n'.join([m.strip() for m in think_matches])
                content = re.sub(r'<think>.*?</think>', '', content, flags=re.DOTALL).strip()

            final_answer = content
    except Exception as e:
        final_answer = f"❌ 生成回答失败：{str(e)}"
        thinking_content = ""

    return {
        "question": original_question,
        "rewritten_query": rewritten_query,
        "local_knowledge": all_local_knowledge,
        "used_knowledge": used_knowledge,
        "online_knowledge": online_knowledge,
        "web_knowledge": web_knowledge,
        "final_answer": final_answer,
        "thinking": thinking_content
    }


# ===================== 历史记录管理 =====================
def update_conversation_history(conv_id: str, question: str, answer: str, thinking: str = "") -> bool:
    """更新对话历史"""
    from .utils import get_all_conversations, save_all_conversations
    import uuid

    conversations = get_all_conversations()
    conv = next((c for c in conversations if c.get('id') == conv_id), None)

    if not conv:
        return False

    conv['messages'].insert(0, {
        "question": question,
        "answer": answer,
        "thinking": thinking,
        "time": time.strftime("%Y-%m-%d %H:%M:%S")
    })

    if len(conv['messages']) == 1:
        conv['title'] = question[:30] + ('...' if len(question) > 30 else '')

    conv['time'] = time.strftime("%Y-%m-%d %H:%M:%S")
    save_all_conversations(conversations)
    return True